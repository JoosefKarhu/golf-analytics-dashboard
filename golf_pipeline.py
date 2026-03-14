#!/usr/bin/env python3
"""
Golf Pipeline v1.0
──────────────────────────────────────────────────────────────────────────────
Automated processor for Virtual Golf 3 simulator screenshots.

What it does:
  1. Scans the Golf folder for new image files
  2. Sends them to Claude AI for vision analysis
  3. Stores extracted round data in rounds_data.json
  4. Regenerates Joonas_Golf_Tracker.xlsx
  5. Regenerates Joonas_Golf_Analysis_HCP5.docx
  6. Archives processed images to the "Round data" subfolder

Requirements:
    pip install openpyxl python-docx

Usage:
    export ANTHROPIC_API_KEY=sk-ant-...
    python golf_pipeline.py

Options:
    --folder PATH     Golf folder path (default: folder containing this script)
    --api-key KEY     Anthropic API key (default: $ANTHROPIC_API_KEY env var)
    --dry-run         Analyse images and print extracted data without saving
    --seed            Initialise the data store with the existing 14 rounds
                      (run once if rounds_data.json doesn't exist yet)
──────────────────────────────────────────────────────────────────────────────
"""

import os, sys, json, base64, shutil, argparse, zipfile, re, urllib.request, urllib.error
from pathlib import Path
from datetime import datetime

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
EXCLUDE_FILES    = {"Joonas_Golf_Tracker.xlsx", "Joonas_Golf_Analysis_HCP5.docx",
                    "golf_pipeline.py", "rounds_data.json"}
CLAUDE_MODEL     = "claude-sonnet-4-6"
API_URL          = "https://api.anthropic.com/v1/messages"
MAX_IMAGES_BATCH = 40   # images per API call (keeps tokens manageable)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DATA STORE  (rounds_data.json)
# ══════════════════════════════════════════════════════════════════════════════

def load_store(store_path: Path) -> dict:
    if store_path.exists():
        with open(store_path) as f:
            return json.load(f)
    return {"rounds": [], "range_sessions": [], "processed_images": []}


def save_store(data: dict, store_path: Path):
    with open(store_path, "w") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — IMAGE SCANNER
# ══════════════════════════════════════════════════════════════════════════════

def find_new_images(folder: Path, store: dict) -> list[Path]:
    """Return image files in folder root that haven't been processed yet."""
    already_done = set(store.get("processed_images", []))
    images = []
    for f in sorted(folder.iterdir()):
        if (f.is_file()
                and f.suffix.lower() in IMAGE_EXTENSIONS
                and f.name not in EXCLUDE_FILES
                and f.name not in already_done):
            images.append(f)
    return images


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — CLAUDE VISION ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

ANALYSIS_PROMPT = """
You are analysing screenshots from the Virtual Golf 3 golf simulator app.
The images may contain hole-by-hole shot maps and/or round summary/scorecard screens.

Examine ALL provided images carefully and return a single JSON object with this exact schema:

{
  "rounds": [
    {
      "course": "string — course name exactly as shown",
      "date": "YYYY-MM-DD or null if not visible",
      "score_vs_par": integer (e.g. 7 for +7, -1 for -1, 0 for E),
      "total_strokes": integer,
      "par": integer (total course par, derive from strokes - score_vs_par),
      "holes_played": integer (18 or 9 or partial),
      "fairways_hit_pct": integer or null,
      "avg_drive_m": integer or null,
      "longest_drive_m": integer or null,
      "gir_pct": integer or null,
      "scrambling_pct": integer or null,
      "total_putts": integer or null,
      "putts_per_hole": float or null,
      "holes": [
        {
          "hole": integer,
          "par": integer,
          "distance_m": integer,
          "hcp": integer or null,
          "strokes": integer,
          "drive_m": integer or null,
          "approach_location": "fairway|rough|sand|green|deep_rough|ob|unknown",
          "from_pin_m": float or null,
          "putts": integer or null,
          "gir": boolean,
          "clubs_used": ["Dr","3w","5i",etc] or []
        }
      ],
      "source_images": ["filename1.png", "filename2.png", ...]
    }
  ],
  "range_sessions": [
    {
      "date": "YYYY-MM-DD or null",
      "club": "Driver|3w|7i|etc",
      "shots": integer or null,
      "target_hit_pct": integer or null,
      "avg_carry_m": integer or null,
      "longest_m": integer or null,
      "avg_from_pin_m": float or null,
      "source_images": ["filename.png"]
    }
  ]
}

Rules:
- Group hole-map images with their corresponding round summary into ONE round entry.
- If you see two different courses, create two separate round entries.
- Infer par = total_strokes - score_vs_par.
- For holes: gir=true if the approach shot lands on the green in regulation.
- Return ONLY the JSON object — no explanation, no markdown fences.
- If a field is genuinely not visible, use null.
"""


def encode_image(path: Path) -> dict:
    """Encode an image file as a base64 content block for the API."""
    suffix = path.suffix.lower().lstrip(".")
    mime   = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png"}.get(suffix, "image/png")
    with open(path, "rb") as f:
        data = base64.standard_b64encode(f.read()).decode()
    return {"type": "image", "source": {"type": "base64", "media_type": mime, "data": data}}


def call_claude(api_key: str, images: list[Path]) -> dict:
    """Send images to Claude and return the parsed JSON response."""
    content = []
    for img in images:
        content.append(encode_image(img))
        content.append({"type": "text", "text": f"[Image filename: {img.name}]"})
    content.append({"type": "text", "text": ANALYSIS_PROMPT})

    payload = json.dumps({
        "model": CLAUDE_MODEL,
        "max_tokens": 8192,
        "messages": [{"role": "user", "content": content}]
    }).encode()

    req = urllib.request.Request(
        API_URL,
        data=payload,
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode()
        raise RuntimeError(f"API error {e.code}: {body}") from e

    text = result["content"][0]["text"].strip()
    # Strip markdown code fences if present
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def analyse_images(api_key: str, images: list[Path]) -> dict:
    """Analyse images in batches and merge results."""
    merged = {"rounds": [], "range_sessions": []}
    for i in range(0, len(images), MAX_IMAGES_BATCH):
        batch = images[i:i + MAX_IMAGES_BATCH]
        print(f"  Sending batch {i//MAX_IMAGES_BATCH + 1} "
              f"({len(batch)} images) to Claude…")
        result = call_claude(api_key, batch)
        merged["rounds"].extend(result.get("rounds", []))
        merged["range_sessions"].extend(result.get("range_sessions", []))
    return merged


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — XLSX GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

# ── Colour/style constants ────────────────────────────────────────────────────
S_GREEN  = PatternFill("solid", fgColor="1A6B3A")
S_GOLD   = PatternFill("solid", fgColor="C9A84C")
S_ALT    = PatternFill("solid", fgColor="F2F2F2")
S_RED_F  = PatternFill("solid", fgColor="CC0000")
S_NONE   = PatternFill(fill_type=None)

F_WHITE  = Font(name="Arial", bold=True,  color="FFFFFF", size=10)
F_DARK   = Font(name="Arial",             color="1A1A1A", size=10)
F_GREEN  = Font(name="Arial",             color="1A6B3A", size=10)
F_RED    = Font(name="Arial",             color="CC0000", size=10)
F_GOLD_B = Font(name="Arial", bold=True,  color="C9A84C", size=10)
F_GREY   = Font(name="Arial",             color="888888", size=9)
CENTER   = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT     = Alignment(horizontal="left",   vertical="center", wrap_text=True)
_thin    = Side(style="thin", color="CCCCCC")
BORDER   = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)


def _cell(ws, row, col, value=None, font=None, fill=None, align=CENTER, border=BORDER):
    c = ws.cell(row=row, column=col, value=value)
    if font:   c.font      = font
    if fill:   c.fill      = fill
    if align:  c.alignment = align
    if border: c.border    = border
    return c


def generate_xlsx(store: dict, out_path: Path):
    """Regenerate the full workbook from the data store."""
    wb = openpyxl.Workbook()

    rounds = [r for r in store["rounds"] if r.get("holes_played", 18) >= 18]
    rounds_all = store["rounds"]   # includes partial
    range_sessions = store.get("range_sessions", [])

    # ── Sheet 1: Round History ────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Round History"

    # Title rows
    ws.merge_cells("A1:N1")
    _cell(ws, 1, 1, "⛳  JOONAS – GOLF PERFORMANCE TRACKER  |  Journey to HCP 5",
          font=Font(name="Arial", bold=True, color="FFFFFF", size=13),
          fill=S_GREEN, align=CENTER)
    ws.row_dimensions[1].height = 28

    ws.merge_cells("A2:N2")
    _cell(ws, 2, 1, "⚠  All rounds played from Championship (Back) Tees — metrics reflect maximum course difficulty",
          font=Font(name="Arial", bold=True, color="FFFFFF", size=10),
          fill=S_RED_F, align=CENTER)
    ws.row_dimensions[2].height = 20

    ws.merge_cells("A3:N3")
    _cell(ws, 3, 1,
          f"Virtual Golf 3 Simulator Data  •  All distances in metres  •  Updated {datetime.now().strftime('%b %d, %Y')}",
          font=F_GREY, fill=S_NONE, align=CENTER, border=None)
    ws.row_dimensions[3].height = 16

    # Header row
    headers = ["#","Date","Course","Type","Score\n(+/-)","Holes","Strokes",
               "Fairways\nHit %","Avg Drive\n(m)","Longest\nDrive (m)",
               "GIR %","Scrambling\n%","Putts/Hole\n(Avg)","Total\nPutts"]
    for col, h in enumerate(headers, 1):
        _cell(ws, 4, col, h, font=F_WHITE, fill=S_GREEN)
    ws.row_dimensions[4].height = 32

    # Data rows
    for idx, rd in enumerate(rounds_all, 1):
        row = 4 + idx
        fill = S_ALT if idx % 2 == 0 else S_NONE
        vals = [
            idx,
            rd.get("date", ""),
            rd.get("course", ""),
            rd.get("type", "Full"),
            rd.get("score_vs_par"),
            rd.get("holes_played", 18),
            rd.get("total_strokes"),
            rd.get("fairways_hit_pct"),
            rd.get("avg_drive_m"),
            rd.get("longest_drive_m"),
            rd.get("gir_pct"),
            rd.get("scrambling_pct"),
            rd.get("putts_per_hole"),
            rd.get("total_putts"),
        ]
        for col, v in enumerate(vals, 1):
            c = _cell(ws, row, col, v, font=F_DARK, fill=fill)
        # Colour score
        score = rd.get("score_vs_par")
        gir   = rd.get("gir_pct")
        if score is not None:
            ws.cell(row, 5).font = F_GREEN if score <= 2 else (F_RED if score >= 7 else F_DARK)
        if gir is not None:
            ws.cell(row, 11).font = F_GREEN if gir >= 60 else (F_RED if gir <= 30 else F_DARK)

    # AVG row
    avg_row = 4 + len(rounds_all) + 1
    dr = 4 + len(rounds_all)           # last data row
    ws.merge_cells(f"A{avg_row}:D{avg_row}")
    _cell(ws, avg_row, 1, "AVG", font=F_GOLD_B, fill=S_GOLD)
    avg_cols = {5:"E",7:"G",8:"H",9:"I",10:"J",11:"K",12:"L",13:"M"}
    for col, letter in avg_cols.items():
        _cell(ws, avg_row, col, f"=AVERAGE({letter}5:{letter}{dr})",
              font=F_GOLD_B, fill=S_GOLD)

    # Column widths
    widths = [4, 13, 28, 7, 8, 7, 9, 10, 10, 12, 7, 12, 12, 10]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Sheet 2: Range Sessions ───────────────────────────────────────────────
    ws2 = wb.create_sheet("Range Sessions")
    ws2.merge_cells("A1:G1")
    _cell(ws2, 1, 1, "🎯  RANGE SESSION LOG", font=Font(name="Arial", bold=True, color="FFFFFF", size=12),
          fill=S_GREEN, align=CENTER)
    rs_headers = ["Date","Club","Shots","Target Hit %","Avg Carry (m)","Longest (m)","Avg From Pin (m)"]
    for col, h in enumerate(rs_headers, 1):
        _cell(ws2, 2, col, h, font=F_WHITE, fill=S_GREEN)
    for idx, rs in enumerate(range_sessions, 1):
        row = 2 + idx
        fill = S_ALT if idx % 2 == 0 else S_NONE
        for col, v in enumerate([
            rs.get("date",""), rs.get("club",""), rs.get("shots"),
            rs.get("target_hit_pct"), rs.get("avg_carry_m"),
            rs.get("longest_m"), rs.get("avg_from_pin_m")
        ], 1):
            _cell(ws2, row, col, v, font=F_DARK, fill=fill)
    for i, w in enumerate([13,10,8,14,14,12,16], 1):
        ws2.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Sheet 3: Trends & KPIs ────────────────────────────────────────────────
    ws3 = wb.create_sheet("Trends & KPIs")
    ws3.merge_cells("A1:H1")
    _cell(ws3, 1, 1, "⚠  All rounds from Championship Tees — HCP 5 benchmarks adjusted accordingly",
          font=F_WHITE, fill=S_RED_F, align=CENTER)
    ws3.merge_cells("A2:H2")
    _cell(ws3, 2, 1, "📊  PERFORMANCE TRENDS & KPI DASHBOARD",
          font=Font(name="Arial", bold=True, color="1A6B3A", size=12), align=CENTER, border=None)

    ws3.merge_cells("A4:H4")
    _cell(ws3, 4, 1, "KEY METRICS – Full 18-Hole Rounds vs HCP 5 Benchmark (Championship Tees)",
          font=Font(name="Arial", bold=True, color="1A1A1A", size=10), align=LEFT, border=None)
    kpi_headers = ["Metric","Best","Worst","Avg (All)","HCP 5 Target","Gap","Status"]
    for col, h in enumerate(kpi_headers, 1):
        _cell(ws3, 5, col, h, font=F_WHITE, fill=S_GREEN)

    def _avg(lst): v=[x for x in lst if x is not None]; return round(sum(v)/len(v),1) if v else None
    def _best(lst, asc=True): v=[x for x in lst if x is not None]; return (min(v) if asc else max(v)) if v else None
    def _worst(lst, asc=True): v=[x for x in lst if x is not None]; return (max(v) if asc else min(v)) if v else None

    kpi_data = [
        ("Score vs Par",   _best(  [r["score_vs_par"]   for r in rounds], True),
                           _worst( [r["score_vs_par"]   for r in rounds], True),
                           f"+{_avg([r['score_vs_par']  for r in rounds]):.2f}" if _avg([r["score_vs_par"] for r in rounds]) else "—",
                           "≤ +3 *", "~2–3 shots", "🔴 In Progress"),
        ("Strokes (18H)",  _best(  [r["total_strokes"]  for r in rounds], True),
                           _worst( [r["total_strokes"]  for r in rounds], True),
                           f"{_avg([r['total_strokes']  for r in rounds]):.1f}" if _avg([r["total_strokes"] for r in rounds]) else "—",
                           "≤ 75 *", "~2 strokes", "🔴 In Progress"),
        ("Fairways Hit %", _best(  [r["fairways_hit_pct"] for r in rounds if r.get("fairways_hit_pct")], False),
                           _worst( [r["fairways_hit_pct"] for r in rounds if r.get("fairways_hit_pct")], False),
                           f"{_avg([r['fairways_hit_pct'] for r in rounds if r.get('fairways_hit_pct')]):.0f}%" if _avg([r.get("fairways_hit_pct") for r in rounds]) else "—",
                           "≥ 60%", "+5–10%", "🔶 Close"),
        ("GIR %",          _best(  [r["gir_pct"] for r in rounds if r.get("gir_pct")], False),
                           _worst( [r["gir_pct"] for r in rounds if r.get("gir_pct")], False),
                           f"{_avg([r['gir_pct'] for r in rounds if r.get('gir_pct')]):.0f}%" if _avg([r.get("gir_pct") for r in rounds]) else "—",
                           "≥ 58% *", "+10%", "🔴 Needs Work"),
        ("Scrambling %",   _best(  [r["scrambling_pct"] for r in rounds if r.get("scrambling_pct")], False),
                           _worst( [r["scrambling_pct"] for r in rounds if r.get("scrambling_pct")], False),
                           f"{_avg([r['scrambling_pct'] for r in rounds if r.get('scrambling_pct')]):.0f}%" if _avg([r.get("scrambling_pct") for r in rounds]) else "—",
                           "≥ 55%", "+15%", "🔴 Needs Work"),
        ("Putts/Hole",     _best(  [r["putts_per_hole"] for r in rounds if r.get("putts_per_hole")], True),
                           _worst( [r["putts_per_hole"] for r in rounds if r.get("putts_per_hole")], True),
                           f"{_avg([r['putts_per_hole'] for r in rounds if r.get('putts_per_hole')]):.2f}" if _avg([r.get("putts_per_hole") for r in rounds]) else "—",
                           "≤ 1.70", "–0.03", "🔶 Close"),
    ]
    for i, row_data in enumerate(kpi_data):
        row = 6 + i
        fill = S_ALT if i % 2 else S_NONE
        for col, val in enumerate(row_data, 1):
            _cell(ws3, row, col, val, font=F_DARK, fill=fill)

    # Score trend table
    ws3.merge_cells("A13:H13")
    _cell(ws3, 13, 1, "SCORE TREND — Chronological (18-Hole Full Rounds)",
          font=Font(name="Arial", bold=True, color="1A1A1A", size=10), align=LEFT, border=None)
    for col, h in enumerate(["Round","Date","Course","Score","Strokes","GIR %","FW %"], 1):
        _cell(ws3, 14, col, h, font=F_WHITE, fill=S_GREEN)
    for idx, rd in enumerate(rounds, 1):
        row = 14 + idx
        fill = S_ALT if idx % 2 else S_NONE
        for col, v in enumerate([idx, rd.get("date",""), rd.get("course",""),
                                  rd.get("score_vs_par"), rd.get("total_strokes"),
                                  rd.get("gir_pct"), rd.get("fairways_hit_pct")], 1):
            _cell(ws3, row, col, v, font=F_DARK, fill=fill)
    for i, w in enumerate([7,13,28,8,10,8,8], 1):
        ws3.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Sheet 4: HCP 5 Roadmap ────────────────────────────────────────────────
    ws4 = wb.create_sheet("HCP 5 Roadmap")
    ws4.merge_cells("A1:E1")
    _cell(ws4, 1, 1, "🏌  HCP 5 ROADMAP  |  Key Milestones & Practice Targets",
          font=Font(name="Arial", bold=True, color="FFFFFF", size=12), fill=S_GREEN, align=CENTER)
    milestones = [
        ("Milestone","Current","Target","Priority","Timeline"),
        ("GIR % (champ tees)", f"{_avg([r.get('gir_pct') for r in rounds if r.get('gir_pct')]):.0f}%",
         "58%+", "🔴 #1", "2–3 months"),
        ("Fairways Hit %", f"{_avg([r.get('fairways_hit_pct') for r in rounds if r.get('fairways_hit_pct')]):.0f}%",
         "60%+", "🔴 #2", "2–3 months"),
        ("Scrambling %", f"{_avg([r.get('scrambling_pct') for r in rounds if r.get('scrambling_pct')]):.0f}%",
         "55%+", "🔴 #3", "3–4 months"),
        ("Putts/Hole", f"{_avg([r.get('putts_per_hole') for r in rounds if r.get('putts_per_hole')]):.2f}",
         "≤ 1.70", "🔶 #4", "Ongoing"),
        ("Avg Score vs Par", f"+{_avg([r['score_vs_par'] for r in rounds]):.1f}",
         "+3 or better", "🎯 Goal", "Aug–Oct 2026"),
        ("Par 3 GIR %", "~35% est.", "50%+", "🔴 Urgent", "1–2 months"),
        ("Wedge from 40–80m", "Inconsistent", "≤ 8m from pin", "🔴 Urgent", "1–2 months"),
    ]
    for i, row_data in enumerate(milestones):
        row = 2 + i
        is_hdr = i == 0
        fill = S_GREEN if is_hdr else (S_ALT if i % 2 == 0 else S_NONE)
        font = F_WHITE if is_hdr else F_DARK
        for col, val in enumerate(row_data, 1):
            _cell(ws4, row, col, val, font=font, fill=fill)
    for i, w in enumerate([28, 14, 14, 12, 16], 1):
        ws4.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Sheet 5: Club Distances ───────────────────────────────────────────────
    ws5 = wb.create_sheet("Club Distances")
    ws5.merge_cells("A1:D1")
    _cell(ws5, 1, 1, "⛳  CLUB DISTANCE REFERENCE  (from range session data)",
          font=Font(name="Arial", bold=True, color="FFFFFF", size=11), fill=S_GREEN, align=CENTER)
    for col, h in enumerate(["Club","Avg Carry (m)","Longest (m)","Target Hit %"], 1):
        _cell(ws5, 2, col, h, font=F_WHITE, fill=S_GREEN)
    club_data = [
        ("Driver",   226, 304, "33%"),
        ("3-Wood",   206, None,"64%"),
        ("1-Iron",   None,None,None),
        ("3-Iron",   None,None,None),
        ("4-Iron",   None,None,None),
        ("5-Iron",   None,None,None),
        ("6-Iron",   None,None,None),
        ("7-Iron",   152, None,"42%"),
        ("8-Iron",   134, None,None),
        ("9-Iron",   120, None,None),
        ("PW",       100, None,None),
        ("56° Wedge", 55, None,None),
        ("52° Wedge", 75, None,None),
    ]
    for idx, row_data in enumerate(club_data, 1):
        fill = S_ALT if idx % 2 == 0 else S_NONE
        for col, v in enumerate(row_data, 1):
            _cell(ws5, 2 + idx, col, v, font=F_DARK, fill=fill)
    for i, w in enumerate([14, 16, 14, 14], 1):
        ws5.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    wb.save(out_path)
    print(f"  ✓ xlsx saved: {out_path.name}")


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DOCX GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

C_GREEN = RGBColor(0x1A, 0x6B, 0x3A)
C_GOLD  = RGBColor(0xC9, 0xA8, 0x4C)
C_RED   = RGBColor(0xCC, 0x00, 0x00)
C_DARK  = RGBColor(0x1A, 0x1A, 0x1A)
C_MID   = RGBColor(0x66, 0x66, 0x66)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_GREEN = "1A6B3A"; HEX_GOLD = "C9A84C"; HEX_RED = "CC0000"
HEX_LGREY = "F2F2F2"; HEX_WHITE = "FFFFFF"


def _hex_rgb(val):
    if isinstance(val, RGBColor): return val
    h = val.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def _shd(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def _cp(cell, text, bold=False, color=C_DARK, size=10, align="c", italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align=="c" else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.color.rgb=_hex_rgb(color) if not isinstance(color, RGBColor) else color


def _heading(doc, text, level=1, color=C_GREEN):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.bold=True; r.font.color.rgb=color
    sizes = {1:18, 2:13, 3:11}
    r.font.size = Pt(sizes.get(level,11))
    before = {1:18, 2:12, 3:8}
    p.paragraph_format.space_before = Pt(before.get(level,8))
    p.paragraph_format.space_after  = Pt(6 if level==1 else 4)
    return p


def _body(doc, text, color=C_DARK):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.size=Pt(10); r.font.color.rgb=color
    r.font.name="Arial"; return p


def _bullet(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    if prefix:
        r1 = p.add_run(prefix+" "); r1.bold=True; r1.font.size=Pt(10); r1.font.color.rgb=C_GREEN
    r2 = p.add_run(text); r2.font.size=Pt(10); r2.font.color.rgb=C_DARK


def _divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"6")
    bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),HEX_GREEN)
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


def _table(doc, rows, col_widths=None):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    for i, row_data in enumerate(rows):
        for j, (text, bold, bg, color, align) in enumerate(row_data):
            cell = tbl.cell(i, j)
            if bg: _shd(cell, bg)
            _cp(cell, text, bold=bold, color=color or C_DARK, align=align)
    return tbl


def _avg_r(lst):
    v = [x for x in lst if x is not None]
    return round(sum(v)/len(v), 2) if v else None


def generate_docx(store: dict, out_path: Path):
    """Generate a full performance report from the data store."""
    doc = Document()
    section = doc.sections[0]
    section.page_width=Inches(8.27); section.page_height=Inches(11.69)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(section, attr, Inches(1.0))

    rounds_all = store["rounds"]
    rounds     = [r for r in rounds_all if r.get("holes_played",18) >= 18]
    n          = len(rounds)
    courses    = len({r["course"] for r in rounds_all})

    # ── Cover ─────────────────────────────────────────────────────────────────
    for text, size, color, bold, italic in [
        ("⛳  GOLF PERFORMANCE ANALYSIS", 24, C_GREEN, True,  False),
        ("JOONAS  |  Journey to Handicap 5", 16, C_GOLD, True, False),
        (f"Virtual Golf 3 Simulator Data  •  Nov 2025 – Mar 2026", 11, C_MID, False, False),
        ("⚠  All rounds played from Championship Tees", 11, C_RED, True, False),
        ("Prepared with Claude AI  |  March 2026", 10, C_MID, False, True),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text == "⛳  GOLF PERFORMANCE ANALYSIS":
            p.paragraph_format.space_before = Pt(24)
        r = p.add_run(text); r.bold=bold; r.italic=italic
        r.font.size=Pt(size); r.font.color.rgb=color; r.font.name="Arial"
    _divider(doc); doc.add_paragraph()

    # ── Section 1: Executive Summary ─────────────────────────────────────────
    _heading(doc, "1.  Executive Summary")
    _body(doc, f"This report covers {len(rounds_all)} rounds of course play across {courses} unique "
               f"courses plus range sessions logged in Virtual Golf 3. Goal: Handicap 5.")
    _body(doc, "IMPORTANT: Every round was played from Championship (back) tees. This adds ~400–600m "
               "over regular tees. HCP 5 benchmarks below are adjusted for these conditions.", color=C_RED)
    doc.add_paragraph()

    # KPI table
    scores  = [r["score_vs_par"]    for r in rounds]
    strokes = [r["total_strokes"]   for r in rounds]
    fw_list = [r.get("fairways_hit_pct") for r in rounds if r.get("fairways_hit_pct")]
    gir_l   = [r.get("gir_pct")     for r in rounds if r.get("gir_pct")]
    scr_l   = [r.get("scrambling_pct") for r in rounds if r.get("scrambling_pct")]
    put_l   = [r.get("putts_per_hole") for r in rounds if r.get("putts_per_hole")]
    drv_l   = [r.get("avg_drive_m") for r in rounds if r.get("avg_drive_m")]

    avg_sc  = _avg_r(scores)
    avg_st  = _avg_r(strokes)
    avg_fw  = _avg_r(fw_list)
    avg_gi  = _avg_r(gir_l)
    avg_sc2 = _avg_r(scr_l)
    avg_pu  = _avg_r(put_l)
    avg_dr  = _avg_r(drv_l)

    def _status_score(v, good, bad):
        if v is None: return ("—", C_MID)
        if v <= good: return ("✅ On Target", C_GREEN)
        if v >= bad:  return ("🔴 Gap", C_RED)
        return ("🔶 Close", C_GOLD)

    kpi_rows = [
        [("METRIC",True,HEX_GREEN,HEX_WHITE,"c"),("CURRENT AVG",True,HEX_GREEN,HEX_WHITE,"c"),
         ("HCP 5 TARGET",True,HEX_GREEN,HEX_WHITE,"c"),("STATUS",True,HEX_GREEN,HEX_WHITE,"c")],
        [("Score vs Par",False,None,None,"c"), (f"+{avg_sc:.2f}" if avg_sc else "—",False,None,None,"c"),
         ("≤ +3 *",False,None,None,"c"), ("🔴 In Progress",False,None,HEX_RED,"c")],
        [("Strokes (18H)",False,HEX_LGREY,None,"c"),(f"{avg_st:.1f}" if avg_st else "—",False,HEX_LGREY,None,"c"),
         ("≤ 75 *",False,HEX_LGREY,None,"c"),("🔴 In Progress",False,HEX_LGREY,HEX_RED,"c")],
        [("Fairways Hit",False,None,None,"c"),(f"{avg_fw:.0f}%" if avg_fw else "—",False,None,None,"c"),
         ("≥ 60%",False,None,None,"c"),
         ("🔶 Close" if avg_fw and avg_fw >= 55 else "🔴 Gap",False,None,
          HEX_GOLD if avg_fw and avg_fw >= 55 else HEX_RED,"c")],
        [("GIR %",False,HEX_LGREY,None,"c"),(f"{avg_gi:.0f}%" if avg_gi else "—",False,HEX_LGREY,None,"c"),
         ("≥ 58% *",False,HEX_LGREY,None,"c"),("🔴 Gap",False,HEX_LGREY,HEX_RED,"c")],
        [("Scrambling",False,None,None,"c"),(f"{avg_sc2:.0f}%" if avg_sc2 else "—",False,None,None,"c"),
         ("≥ 55%",False,None,None,"c"),("🔴 Gap",False,None,HEX_RED,"c")],
        [("Putts/Hole",False,HEX_LGREY,None,"c"),(f"{avg_pu:.2f}" if avg_pu else "—",False,HEX_LGREY,None,"c"),
         ("≤ 1.70",False,HEX_LGREY,None,"c"),
         ("🔶 Close" if avg_pu and avg_pu <= 1.75 else "🔴 Gap",False,HEX_LGREY,
          HEX_GOLD if avg_pu and avg_pu <= 1.75 else HEX_RED,"c")],
        [("Avg Drive",False,None,None,"c"),(f"{avg_dr:.0f}m" if avg_dr else "—",False,None,None,"c"),
         ("≥ 240m",False,None,None,"c"),
         ("🔶 Close" if avg_dr and avg_dr >= 225 else "🔴 Gap",False,None,
          HEX_GOLD if avg_dr and avg_dr >= 225 else HEX_RED,"c")],
    ]
    _table(doc, kpi_rows, col_widths=[2.2, 1.5, 1.5, 1.5])
    doc.add_paragraph()

    # Find best round
    best = min(rounds, key=lambda r: r["score_vs_par"])
    _body(doc, f"* Benchmarks adjusted for Championship Tee conditions (standard HCP 5 = ≤+2 from regular tees).", color=C_MID)
    doc.add_paragraph()
    _body(doc, f"Best round: {best['course']} on {best.get('date','')} — "
               f"+{best['score_vs_par']} ({best['total_strokes']} strokes), "
               f"{best.get('gir_pct','—')}% GIR, {best.get('fairways_hit_pct','—')}% FW. "
               f"This confirms HCP 5 ball-striking capability already exists.")
    _divider(doc)

    # ── Section 2: Round-by-Round ─────────────────────────────────────────────
    _heading(doc, "2.  Round-by-Round Results")
    _body(doc, "All 18-hole and notable partial rounds. Distances in metres. Championship tees throughout.")
    doc.add_paragraph()

    hdr = [("#",True,HEX_GREEN,HEX_WHITE,"c"),("Date",True,HEX_GREEN,HEX_WHITE,"c"),
           ("Course",True,HEX_GREEN,HEX_WHITE,"c"),("Score",True,HEX_GREEN,HEX_WHITE,"c"),
           ("Stk",True,HEX_GREEN,HEX_WHITE,"c"),("FW%",True,HEX_GREEN,HEX_WHITE,"c"),
           ("GIR%",True,HEX_GREEN,HEX_WHITE,"c"),("Scr%",True,HEX_GREEN,HEX_WHITE,"c"),
           ("Putts",True,HEX_GREEN,HEX_WHITE,"c")]
    body_rows = []
    for idx, rd in enumerate(rounds_all, 1):
        fill = HEX_LGREY if idx % 2 == 0 else None
        sc   = rd.get("score_vs_par")
        gi   = rd.get("gir_pct")
        sc_col = HEX_GREEN if sc is not None and sc<=2 else (HEX_RED if sc is not None and sc>=7 else None)
        gi_col = HEX_GREEN if gi is not None and gi>=60 else (HEX_RED if gi is not None and gi<=30 else None)
        body_rows.append([
            (str(idx),False,fill,None,"c"),
            (rd.get("date","")[-5:].replace("-","/") if rd.get("date") else "—",False,fill,None,"c"),
            (rd.get("course","")[:26],False,fill,None,"l"),
            (f"+{sc}" if sc and sc>0 else str(sc) if sc is not None else "—",False,fill,sc_col,"c"),
            (str(rd.get("total_strokes","—")),False,fill,None,"c"),
            (f"{rd.get('fairways_hit_pct','—')}%" if rd.get("fairways_hit_pct") else "—",False,fill,None,"c"),
            (f"{gi}%" if gi else "—",False,fill,gi_col,"c"),
            (f"{rd.get('scrambling_pct','—')}%" if rd.get("scrambling_pct") else "—",False,fill,None,"c"),
            (str(rd.get("total_putts","—")),False,fill,None,"c"),
        ])
    _table(doc, [hdr]+body_rows, col_widths=[0.3,1.0,2.2,0.65,0.5,0.65,0.65,0.65,0.6])
    doc.add_paragraph()
    _divider(doc)

    # ── Section 3: Key Findings ───────────────────────────────────────────────
    _heading(doc, "3.  Key Findings & Analysis")

    # 3.1 Scoring trend
    _heading(doc, "3.1  Scoring Trend", level=2)
    best_r  = min(rounds, key=lambda r: r["score_vs_par"])
    worst_r = max(rounds, key=lambda r: r["score_vs_par"])
    recent5 = rounds[-5:] if len(rounds) >= 5 else rounds
    r5_avg  = _avg_r([r["score_vs_par"] for r in recent5])
    _body(doc, f"Across {n} full rounds, average score is +{avg_sc:.2f} (range: "
               f"+{min(scores)} to +{max(scores)}). "
               f"Best round: {best_r['course']} (+{best_r['score_vs_par']}, "
               f"{best_r['total_strokes']} strokes, {best_r.get('gir_pct','—')}% GIR). "
               f"Most recent 5 rounds average: +{r5_avg:.1f}. "
               f"The variance between rounds is the central challenge — eliminating blow-up rounds "
               f"(+{worst_r['score_vs_par']} at {worst_r['course']}) will pull the average toward +3.")

    # 3.2 GIR correlation
    _heading(doc, "3.2  GIR — The Scoring Engine", level=2)
    high_gir = [r for r in rounds if r.get("gir_pct") and r["gir_pct"] >= 55]
    low_gir  = [r for r in rounds if r.get("gir_pct") and r["gir_pct"] <= 35]
    _body(doc, f"GIR is the strongest predictor of score in this dataset. Rounds with GIR ≥ 55% "
               f"average +{_avg_r([r['score_vs_par'] for r in high_gir]):.1f}. "
               f"Rounds with GIR ≤ 35% average +{_avg_r([r['score_vs_par'] for r in low_gir]):.1f}. "
               f"Overall GIR average: {avg_gi:.0f}% — target 58%+ from championship tees.")
    for r in sorted(rounds, key=lambda x: x.get("gir_pct") or 0, reverse=True)[:4]:
        _bullet(doc, f"{r['course']}: {r.get('gir_pct','—')}% GIR → +{r['score_vs_par']} ({r['total_strokes']} stk)")

    # 3.3 Scrambling
    _heading(doc, "3.3  Scrambling — Still the Weak Link", level=2)
    _body(doc, f"Scrambling average: {avg_sc2:.0f}% — well below the 55% target. "
               f"When GIR drops, scrambling must compensate; currently it averages only {avg_sc2:.0f}%, "
               f"meaning missed greens almost always result in bogey or worse. "
               f"Best scrambling round: {max(rounds, key=lambda r: r.get('scrambling_pct') or 0).get('course','—')} "
               f"({max([r.get('scrambling_pct') or 0 for r in rounds])}%). "
               f"Improving to 50%+ will act as a floor-raiser on bad ball-striking days.")

    # 3.4 Putting
    _heading(doc, "3.4  Putting — A Genuine Strength", level=2)
    _body(doc, f"Putting average of {avg_pu:.2f}/hole is the closest metric to HCP 5 standard. "
               f"Best putting round: {min(rounds, key=lambda r: r.get('putts_per_hole') or 99).get('course','—')} "
               f"({min([r.get('putts_per_hole') or 99 for r in rounds]):.2f}/hole). "
               f"Home putting mat practice is clearly paying off — continue daily.")

    # 3.5 Course management (data-driven)
    _heading(doc, "3.5  Key Course Management Findings", level=2, color=C_RED)
    _bullet(doc, "Par 3s above 155m are consistently producing bogeys/doubles. Long irons (1i, 3i) "
                 "on par 3s are high-risk — replace with hybrid for 160m+ holes.", "Par 3 Club Selection:")
    _bullet(doc, "Wedge play from 40–80m is leaving the ball 30–38m from the pin on average — "
                 "the root cause of missed GIR-to-score conversion. Dedicated range work at 40m, 50m, 60m required.",
                 "Wedge Distance Control:")
    _bullet(doc, "Hitting fairways does not guarantee GIR: Marcella 92% FW → 44% GIR, Eden 71% FW → 28% GIR. "
                 "Iron approach selection from fairway lies needs focus.", "Fairway-to-GIR Conversion:")
    _bullet(doc, "Catastrophic tee shots (75m drive at Olympic Club H4, 192m into deep rough at Longleaf H5) "
                 "can wreck an otherwise solid round. Use an intermediate target 1m ahead of the ball "
                 "on long, intimidating tee shots.", "Blow-Up Tee Shots:")
    _divider(doc)

    # ── Section 4: Action Plan ────────────────────────────────────────────────
    _heading(doc, "4.  Priority Action Plan — Path to HCP 5")

    _heading(doc, "PRIORITY 1: Iron Accuracy — Hybrid for Par 3s over 155m", level=2, color=C_RED)
    _bullet(doc, "Replace 1-iron and 3-iron on par 3s above 155m with a hybrid. Higher launch, "
                 "softer landing, far better dispersion.", "Immediate Fix:")
    _bullet(doc, "Iron range sessions 3× per week: 7i, 8i, 9i (120–160m range). "
                 "Target: from-pin average below 15m on 7i.", "Weekly Practice:")

    _heading(doc, "PRIORITY 2: Wedge Distance Control (40–80m zone)", level=2, color=C_RED)
    _bullet(doc, "In range sessions, dedicate 20 min to: 40m, 50m, 60m shots with 56°. "
                 "Track where ball finishes. Target: consistently within 8m from 50m.", "Range Drill:")
    _bullet(doc, "On course, commit to a specific number (e.g. 'this is a 52m shot') "
                 "rather than 'somewhere around here'. Precision beats power at this distance.", "Course Mindset:")

    _heading(doc, "PRIORITY 3: Scrambling (current 39% → target 55%)", level=2, color=C_RED)
    _bullet(doc, "30 min per session on pitching/chipping from within 50m. "
                 "Focus on landing zone, not the hole.", "Short Game Practice:")
    _bullet(doc, "When in rough near the green: putt from fringe when possible. "
                 "Fewer moving parts = more reliable.", "On-Course Decision:")

    _heading(doc, "PRIORITY 4: Score Management", level=2, color=C_GOLD)
    _bullet(doc, "Hard cap: double bogey maximum per hole. Accept the double, move on.", "Mental Rule:")
    _bullet(doc, "On long, tight par 4s: intermediate target 1m ahead of ball on tee shot. "
                 "Reduces catastrophic misses significantly.", "Tee Shot Routine:")

    _heading(doc, "PRIORITY 5: Leverage Squash Athletic Background", level=2, color=C_GREEN)
    _body(doc, "National-level squash = elite hand-eye coordination, explosive rotational power, "
               "and competitive composure. The athletic ceiling is very high. "
               "The work is technical, not physical — improvement velocity will be faster than average.")
    _divider(doc)

    # ── Section 5: HCP Projection ─────────────────────────────────────────────
    _heading(doc, "5.  Handicap 5 Projection")
    proj_rows = [
        [("SCENARIO",True,HEX_GREEN,HEX_WHITE,"c"),("CONDITIONS",True,HEX_GREEN,HEX_WHITE,"c"),
         ("TIMELINE",True,HEX_GREEN,HEX_WHITE,"c")],
        [("Optimistic",False,None,HEX_GREEN,"c"),
         ("GIR 58%+ consistently; hybrid on par 3s; wedge control improves to ≤8m from 50m; "
          "squash athleticism accelerates technique gains",False,None,None,"l"),
         ("Aug–Oct 2026",False,None,HEX_GREEN,"c")],
        [("Realistic",False,HEX_LGREY,HEX_GOLD,"c"),
         ("GIR reaches 55%; scrambling to 45–50%; score variance tightens; "
          "par 3 improvement ongoing",False,HEX_LGREY,None,"l"),
         ("Dec 2026–Mar 2027",False,HEX_LGREY,HEX_GOLD,"c")],
        [("Conservative",False,None,HEX_RED,"c"),
         ("GIR stays sub-50%; blow-up holes continue; no structured practice",
          False,None,None,"l"),
         ("18+ months",False,None,HEX_RED,"c")],
    ]
    _table(doc, proj_rows, col_widths=[1.3, 3.8, 2.1])
    doc.add_paragraph()
    _body(doc, f"The best round ({best_r['course']}, +{best_r['score_vs_par']}) proves the ball-striking "
               f"capability is already present. The remaining work is reproducibility — making +1 to +3 "
               f"the norm rather than the exception. With focused practice the optimistic timeline is achievable.")
    _divider(doc)

    # ── Section 6: Data Notes ─────────────────────────────────────────────────
    _heading(doc, "6.  Data Notes & Methodology")
    _body(doc, "Data extracted from Virtual Golf 3 simulator screenshots via Claude AI vision analysis.")
    _bullet(doc, f"{len(rounds_all)} rounds total ({len([r for r in rounds_all if r.get('holes_played',18)<18])} partial, "
                 f"{n} × 18 holes), across {courses} unique courses")
    _bullet(doc, f"Sessions span {rounds_all[0].get('date','?')} to {rounds_all[-1].get('date','?')}")
    _bullet(doc, "All rounds played from Championship (back) tees throughout")
    _bullet(doc, "Some putting/scrambling values auto-assigned by simulator; flagged where noted")

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated by Golf Pipeline v1.0  |  Claude AI  |  {datetime.now().strftime('%B %Y')}")
    r.italic=True; r.font.size=Pt(9); r.font.color.rgb=C_MID

    doc.save(out_path)
    print(f"  ✓ docx saved (pre-fix): {out_path.name}")


def fix_zoom(docx_path: Path):
    """Fix the python-docx zoom attribute bug (bestFit → percent=100)."""
    import tempfile, zipfile as zf
    tmp = Path(tempfile.mktemp(suffix=".docx"))
    with zf.ZipFile(docx_path, "r") as zin, zf.ZipFile(tmp, "w", zf.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                data = data.replace(b'<w:zoom w:val="bestFit"/>',
                                    b'<w:zoom w:percent="100"/>')
            zout.writestr(item, data)
    shutil.move(str(tmp), str(docx_path))
    print(f"  ✓ docx zoom fix applied")


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ARCHIVER
# ══════════════════════════════════════════════════════════════════════════════

def archive_images(images: list[Path], archive_folder: Path):
    archive_folder.mkdir(exist_ok=True)
    for img in images:
        dest = archive_folder / img.name
        shutil.move(str(img), str(dest))
    print(f"  ✓ {len(images)} image(s) archived to '{archive_folder.name}/'")


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SEED DATA  (existing 14 rounds — run once with --seed)
# ══════════════════════════════════════════════════════════════════════════════

SEED_ROUNDS = [
    {"course":"Silverleaf Club","date":"2025-11-18","score_vs_par":0,"total_strokes":72,"par":72,
     "holes_played":18,"fairways_hit_pct":33,"avg_drive_m":224,"longest_drive_m":258,
     "gir_pct":83,"scrambling_pct":None,"putts_per_hole":1.8,"total_putts":33,"holes":[],"source_images":[]},
    {"course":"Sweetens Cove GC","date":"2025-11-18","score_vs_par":5,"total_strokes":41,"par":36,
     "holes_played":9,"fairways_hit_pct":57,"avg_drive_m":201,"longest_drive_m":218,
     "gir_pct":44,"scrambling_pct":20,"putts_per_hole":1.8,"total_putts":16,"holes":[],"source_images":[]},
    {"course":"Bandon Preserve","date":"2025-11-18","score_vs_par":-1,"total_strokes":38,"par":39,
     "holes_played":13,"fairways_hit_pct":None,"avg_drive_m":None,"longest_drive_m":None,
     "gir_pct":77,"scrambling_pct":67,"putts_per_hole":1.7,"total_putts":22,"holes":[],"source_images":[]},
    {"course":"Country Club of Jackson","date":"2025-11-22","score_vs_par":5,"total_strokes":77,"par":72,
     "holes_played":18,"fairways_hit_pct":29,"avg_drive_m":213,"longest_drive_m":238,
     "gir_pct":28,"scrambling_pct":62,"putts_per_hole":None,"total_putts":None,"holes":[],"source_images":[]},
    {"course":"Clear Creek Tahoe","date":"2026-01-20","score_vs_par":3,"total_strokes":74,"par":71,
     "holes_played":18,"fairways_hit_pct":69,"avg_drive_m":233,"longest_drive_m":263,
     "gir_pct":56,"scrambling_pct":38,"putts_per_hole":1.8,"total_putts":32,"holes":[],"source_images":[]},
    {"course":"Emirates GC – Majlis","date":"2026-01-22","score_vs_par":10,"total_strokes":82,"par":72,
     "holes_played":18,"fairways_hit_pct":21,"avg_drive_m":192,"longest_drive_m":236,
     "gir_pct":28,"scrambling_pct":31,"putts_per_hole":None,"total_putts":None,"holes":[],"source_images":[]},
    {"course":"Marcella Club","date":"2026-02-18","score_vs_par":6,"total_strokes":78,"par":72,
     "holes_played":18,"fairways_hit_pct":92,"avg_drive_m":259,"longest_drive_m":287,
     "gir_pct":44,"scrambling_pct":40,"putts_per_hole":1.8,"total_putts":32,"holes":[],"source_images":[]},
    {"course":"Clear Creek Tahoe","date":"2026-02-21","score_vs_par":8,"total_strokes":79,"par":71,
     "holes_played":18,"fairways_hit_pct":69,"avg_drive_m":247,"longest_drive_m":302,
     "gir_pct":50,"scrambling_pct":11,"putts_per_hole":1.8,"total_putts":33,"holes":[],"source_images":[]},
    {"course":"Kawana Hotel Fuji Course","date":"2026-02-24","score_vs_par":2,"total_strokes":74,"par":72,
     "holes_played":18,"fairways_hit_pct":50,"avg_drive_m":250,"longest_drive_m":304,
     "gir_pct":61,"scrambling_pct":57,"putts_per_hole":1.61,"total_putts":29,"holes":[],"source_images":[]},
    {"course":"Olympic Club Lake Course","date":"2026-03-02","score_vs_par":12,"total_strokes":83,"par":71,
     "holes_played":18,"fairways_hit_pct":57,"avg_drive_m":217,"longest_drive_m":259,
     "gir_pct":28,"scrambling_pct":23,"putts_per_hole":1.67,"total_putts":30,"holes":[],"source_images":[]},
    {"course":"Eden Golf Club","date":"2026-03-03","score_vs_par":7,"total_strokes":79,"par":72,
     "holes_played":18,"fairways_hit_pct":71,"avg_drive_m":233,"longest_drive_m":247,
     "gir_pct":28,"scrambling_pct":46,"putts_per_hole":1.56,"total_putts":28,"holes":[],"source_images":[]},
    {"course":"Hualalai Resort","date":"2026-03-03","score_vs_par":6,"total_strokes":78,"par":72,
     "holes_played":18,"fairways_hit_pct":43,"avg_drive_m":229,"longest_drive_m":254,
     "gir_pct":61,"scrambling_pct":43,"putts_per_hole":1.78,"total_putts":32,"holes":[],"source_images":[]},
    {"course":"Cape Wickham Links","date":"2026-03-05","score_vs_par":1,"total_strokes":73,"par":72,
     "holes_played":18,"fairways_hit_pct":86,"avg_drive_m":225,"longest_drive_m":257,
     "gir_pct":78,"scrambling_pct":50,"putts_per_hole":1.83,"total_putts":33,"holes":[],"source_images":[]},
    {"course":"Longleaf Golf & Family Club","date":"2026-03-05","score_vs_par":7,"total_strokes":79,"par":72,
     "holes_played":18,"fairways_hit_pct":71,"avg_drive_m":221,"longest_drive_m":271,
     "gir_pct":44,"scrambling_pct":30,"putts_per_hole":1.56,"total_putts":28,"holes":[],"source_images":[]},
]

SEED_RANGE = [
    {"date":"2026-01-29","club":"Driver","shots":20,"target_hit_pct":14,"avg_carry_m":226,"longest_m":247,"avg_from_pin_m":None,"source_images":[]},
    {"date":"2026-02-16","club":"Driver","shots":5, "target_hit_pct":60,"avg_carry_m":229,"longest_m":241,"avg_from_pin_m":None,"source_images":[]},
    {"date":"2026-02-16","club":"3-Wood","shots":5, "target_hit_pct":64,"avg_carry_m":206,"longest_m":216,"avg_from_pin_m":None,"source_images":[]},
    {"date":"2026-02-16","club":"7-Iron","shots":33,"target_hit_pct":42,"avg_carry_m":152,"longest_m":163,"avg_from_pin_m":None,"source_images":[]},
]


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — MAIN ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Golf Pipeline — automated VG3 data processor")
    parser.add_argument("--folder",  default=str(Path(__file__).parent),
                        help="Path to Golf folder (default: same folder as this script)")
    parser.add_argument("--api-key", default=os.environ.get("ANTHROPIC_API_KEY",""),
                        help="Anthropic API key (default: $ANTHROPIC_API_KEY)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Analyse images and print data without saving")
    parser.add_argument("--seed",    action="store_true",
                        help="Initialise rounds_data.json with the 14 existing rounds (run once)")
    parser.add_argument("--regenerate", action="store_true",
                        help="Skip image analysis and just regenerate xlsx + docx from stored data")
    args = parser.parse_args()

    folder       = Path(args.folder)
    store_path   = folder / "rounds_data.json"
    xlsx_path    = folder / "Joonas_Golf_Tracker.xlsx"
    docx_path    = folder / "Joonas_Golf_Analysis_HCP5.docx"
    archive_dir  = folder / "Round data"

    print(f"\n{'='*60}")
    print(f"  Golf Pipeline v1.0")
    print(f"  Folder : {folder}")
    print(f"{'='*60}\n")

    # ── Seed mode ──────────────────────────────────────────────────────────────
    if args.seed:
        if store_path.exists():
            print("⚠  rounds_data.json already exists. Delete it first to re-seed.")
            sys.exit(1)
        store = {"rounds": SEED_ROUNDS, "range_sessions": SEED_RANGE, "processed_images": []}
        save_store(store, store_path)
        print(f"✓ Seeded {len(SEED_ROUNDS)} rounds + {len(SEED_RANGE)} range sessions → rounds_data.json")
        print("  Now regenerating xlsx and docx from seeded data…")
        generate_xlsx(store, xlsx_path)
        generate_docx(store, docx_path)
        fix_zoom(docx_path)
        print("\n✅ Done. Run without --seed to process new images.")
        return

    # ── Load store ─────────────────────────────────────────────────────────────
    store = load_store(store_path)
    print(f"Loaded data store: {len(store['rounds'])} rounds, "
          f"{len(store.get('range_sessions',[]))} range sessions\n")

    # ── Regenerate-only mode ───────────────────────────────────────────────────
    if args.regenerate:
        print("Regenerating files from stored data (no image analysis)…")
        generate_xlsx(store, xlsx_path)
        generate_docx(store, docx_path)
        fix_zoom(docx_path)
        print("\n✅ Done.")
        return

    # ── Find new images ────────────────────────────────────────────────────────
    new_images = find_new_images(folder, store)
    if not new_images:
        print("No new images found in folder. Nothing to process.")
        print("Tip: drop new screenshots in the Golf folder, then run again.")
        print("     Or use --regenerate to rebuild xlsx/docx from stored data.")
        return

    print(f"Found {len(new_images)} new image(s):")
    for img in new_images:
        print(f"  • {img.name}")
    print()

    if not args.api_key:
        print("❌  No API key provided.")
        print("    Set the ANTHROPIC_API_KEY environment variable or use --api-key KEY")
        sys.exit(1)

    # ── Analyse with Claude ────────────────────────────────────────────────────
    print("Sending images to Claude for analysis…")
    try:
        extracted = analyse_images(args.api_key, new_images)
    except Exception as e:
        print(f"❌  Analysis failed: {e}")
        sys.exit(1)

    print(f"\nExtracted: {len(extracted['rounds'])} round(s), "
          f"{len(extracted.get('range_sessions',[]))} range session(s)\n")

    for rd in extracted["rounds"]:
        print(f"  📍 {rd['course']} | {rd.get('date','?')} | "
              f"+{rd['score_vs_par']} ({rd['total_strokes']} stk) | "
              f"GIR {rd.get('gir_pct','?')}% | FW {rd.get('fairways_hit_pct','?')}%")
    for rs in extracted.get("range_sessions", []):
        print(f"  🎯 Range: {rs['club']} | {rs.get('target_hit_pct','?')}% target hit")

    if args.dry_run:
        print("\n[dry-run] Not saving. Extracted JSON:")
        print(json.dumps(extracted, indent=2))
        return

    # ── Merge into store ───────────────────────────────────────────────────────
    store["rounds"].extend(extracted["rounds"])
    store.setdefault("range_sessions", []).extend(extracted.get("range_sessions", []))
    store.setdefault("processed_images", []).extend([img.name for img in new_images])

    # ── Regenerate files ───────────────────────────────────────────────────────
    print("\nRegenerating xlsx…")
    generate_xlsx(store, xlsx_path)

    print("Regenerating docx…")
    generate_docx(store, docx_path)
    fix_zoom(docx_path)

    # ── Archive images ─────────────────────────────────────────────────────────
    print("\nArchiving images…")
    archive_images(new_images, archive_dir)

    # ── Save store ─────────────────────────────────────────────────────────────
    save_store(store, store_path)
    print(f"  ✓ rounds_data.json updated ({len(store['rounds'])} rounds total)")

    print(f"\n{'='*60}")
    print(f"  ✅  Pipeline complete!")
    print(f"  Rounds in store : {len(store['rounds'])}")
    print(f"  xlsx            : {xlsx_path.name}")
    print(f"  docx            : {docx_path.name}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
